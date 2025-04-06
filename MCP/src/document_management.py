import asyncio
import json
import os
import sys
import base64
import io
import tempfile
from typing import Dict, Any, List, Optional, Union

from fastapi import FastAPI, WebSocket, WebSocketDisconnect
from pydantic import BaseModel, Field

# Path to the main module for importing shared code
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from src.main import app, Message, MessageType, manager

async def handle_parse_document(message_id: str, params: Dict[str, Any], websocket: WebSocket):
    """
    Parse a document and extract text/structure
    Params:
      - file: Base64 encoded file data
      - fileType: Type of file (pdf, docx, txt, etc.)
      - extractImages: Whether to extract images (default: false)
    """
    file_data = params.get("file", "")
    file_type = params.get("fileType", "").lower()
    extract_images = params.get("extractImages", False)
    
    if not file_data or not file_type:
        await manager.send_error(message_id, 400, "File data and type are required", websocket)
        return
    
    try:
        # Create a temporary file
        with tempfile.NamedTemporaryFile(suffix=f".{file_type}", delete=False) as temp_file:
            temp_file_path = temp_file.name
            temp_file.write(base64.b64decode(file_data))
        
        result = {"text": "", "images": [], "metadata": {}}
        
        # Handle different file types
        if file_type == "pdf":
            try:
                import PyPDF2
                
                with open(temp_file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    pages = []
                    
                    # Extract text from each page
                    for page_num in range(len(reader.pages)):
                        page = reader.pages[page_num]
                        pages.append(page.extract_text())
                    
                    # Get document info
                    metadata = reader.metadata
                    if metadata:
                        result["metadata"] = {
                            "title": metadata.get('/Title', ''),
                            "author": metadata.get('/Author', ''),
                            "subject": metadata.get('/Subject', ''),
                            "creator": metadata.get('/Creator', ''),
                            "producer": metadata.get('/Producer', '')
                        }
                    
                    result["text"] = "\n\n".join(pages)
                    result["pageCount"] = len(reader.pages)
                
                # Extract images if requested
                if extract_images:
                    try:
                        from PIL import Image
                        from pdf2image import convert_from_path
                        
                        images = convert_from_path(temp_file_path)
                        result["images"] = []
                        
                        for i, img in enumerate(images):
                            buffered = io.BytesIO()
                            img.save(buffered, format="PNG")
                            img_str = base64.b64encode(buffered.getvalue()).decode()
                            result["images"].append({
                                "page": i + 1,
                                "data": img_str
                            })
                    except ImportError:
                        result["imageExtractionError"] = "pdf2image library not available"
            
            except ImportError:
                result["error"] = "PDF parsing requires PyPDF2"
                result["installCommand"] = "pip install PyPDF2"
        
        elif file_type in ["docx", "doc"]:
            try:
                import docx
                
                doc = docx.Document(temp_file_path)
                paragraphs = [p.text for p in doc.paragraphs]
                result["text"] = "\n\n".join(paragraphs)
                
                # Get document properties
                core_props = doc.core_properties
                result["metadata"] = {
                    "title": core_props.title if hasattr(core_props, 'title') else '',
                    "author": core_props.author if hasattr(core_props, 'author') else '',
                    "comments": core_props.comments if hasattr(core_props, 'comments') else ''
                }
                
                # Extract tables
                tables = []
                for table in doc.tables:
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text for cell in row.cells]
                        table_data.append(row_data)
                    tables.append(table_data)
                
                result["tables"] = tables
            
            except ImportError:
                result["error"] = "DOCX parsing requires python-docx"
                result["installCommand"] = "pip install python-docx"
        
        elif file_type == "txt":
            with open(temp_file_path, 'r', errors='ignore') as file:
                result["text"] = file.read()
        
        elif file_type in ["csv", "tsv"]:
            try:
                import csv
                
                delimiter = ',' if file_type == 'csv' else '\t'
                rows = []
                
                with open(temp_file_path, 'r', newline='', errors='ignore') as file:
                    csv_reader = csv.reader(file, delimiter=delimiter)
                    for row in csv_reader:
                        rows.append(row)
                
                if rows:
                    result["headers"] = rows[0]
                    result["data"] = rows[1:]
                    result["rowCount"] = len(rows) - 1
                    result["columnCount"] = len(rows[0]) if rows[0] else 0
            
            except Exception as e:
                result["error"] = f"Error parsing {file_type.upper()} file: {str(e)}"
        
        # Clean up the temporary file
        try:
            os.unlink(temp_file_path)
        except:
            pass
        
        await manager.send_response(message_id, result, websocket)
    
    except Exception as e:
        await manager.send_error(message_id, 500, str(e), websocket)

async def handle_create_document(message_id: str, params: Dict[str, Any], websocket: WebSocket):
    """
    Create a document from text/markdown content
    Params:
      - content: Document content
      - format: Output format (pdf, docx, html, md)
      - title: Document title
      - template: Optional template to use
    """
    content = params.get("content", "")
    output_format = params.get("format", "pdf").lower()
    title = params.get("title", "Generated Document")
    template = params.get("template", "")
    
    if not content:
        await manager.send_error(message_id, 400, "Document content is required", websocket)
        return
    
    try:
        result = {}
        
        # Generate a temporary file for the output
        output_file = tempfile.NamedTemporaryFile(suffix=f".{output_format}", delete=False)
        output_path = output_file.name
        output_file.close()
        
        # Create a temporary input file with the content
        with tempfile.NamedTemporaryFile(suffix=".md", delete=False) as input_file:
            input_path = input_file.name
            input_file.write(content.encode())
        
        if output_format == "pdf":
            try:
                process = await asyncio.create_subprocess_shell(
                    f"pandoc -f markdown -t pdf --pdf-engine=wkhtmltopdf -o {output_path} {input_path}",
                    stdout=asyncio.subprocess.PIPE,
                    stderr=asyncio.subprocess.PIPE
                )
                stdout, stderr = await process.communicate()
                
                if process.returncode == 0:
                    # Read the generated PDF
                    with open(output_path, 'rb') as file:
                        pdf_data = base64.b64encode(file.read()).decode()
                        result["file"] = pdf_data
                        result["format"] = "pdf"
                else:
                    result["error"] = stderr.decode() if stderr else "Unknown error generating PDF"
                    result["fallbackSuggestion"] = "Install pandoc and wkhtmltopdf for PDF generation"
            except:
                result["error"] = "Could not generate PDF. Make sure pandoc and wkhtmltopdf are installed."
        
        elif output_format == "docx":
            try:
                process = await asyncio.create_subprocess_shell(
                    f"pandoc -f markdown -t docx -o {output_path} {input_path}",
                    stdout=asyncio.subprocess.PIPE,
                    stderr=asyncio.subprocess.PIPE
                )
                stdout, stderr = await process.communicate()
                
                if process.returncode == 0:
                    # Read the generated DOCX
                    with open(output_path, 'rb') as file:
                        docx_data = base64.b64encode(file.read()).decode()
                        result["file"] = docx_data
                        result["format"] = "docx"
                else:
                    result["error"] = stderr.decode() if stderr else "Unknown error generating DOCX"
                    result["fallbackSuggestion"] = "Install pandoc for DOCX generation"
            except:
                result["error"] = "Could not generate DOCX. Make sure pandoc is installed."
        
        elif output_format == "html":
            try:
                import markdown
                
                html_content = markdown.markdown(content)
                html_doc = f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <meta charset="utf-8">
                    <title>{title}</title>
                    <style>
                        body {{ font-family: Arial, sans-serif; line-height: 1.6; max-width: 800px; margin: 0 auto; padding: 20px; }}
                        h1 {{ color: #2c3e50; }}
                        h2 {{ color: #3498db; }}
                        pre {{ background: #f8f8f8; padding: 10px; border-radius: 5px; overflow-x: auto; }}
                        code {{ background: #f0f0f0; padding: 2px 4px; border-radius: 3px; }}
                        blockquote {{ border-left: 4px solid #ccc; padding-left: 15px; color: #555; }}
                    </style>
                </head>
                <body>
                    <h1>{title}</h1>
                    {html_content}
                </body>
                </html>
                """
                
                with open(output_path, 'w') as file:
                    file.write(html_doc)
                
                with open(output_path, 'rb') as file:
                    html_data = base64.b64encode(file.read()).decode()
                    result["file"] = html_data
                    result["format"] = "html"
            
            except ImportError:
                result["error"] = "HTML generation requires markdown package"
                result["installCommand"] = "pip install markdown"
        
        elif output_format == "md":
            # Just return the markdown directly
            with open(output_path, 'w') as file:
                file.write(content)
            
            with open(output_path, 'rb') as file:
                md_data = base64.b64encode(file.read()).decode()
                result["file"] = md_data
                result["format"] = "md"
        
        # Clean up temporary files
        try:
            os.unlink(input_path)
            os.unlink(output_path)
        except:
            pass
        
        await manager.send_response(message_id, result, websocket)
    
    except Exception as e:
        await manager.send_error(message_id, 500, str(e), websocket)

async def handle_convert_document(message_id: str, params: Dict[str, Any], websocket: WebSocket):
    """
    Convert a document from one format to another
    Params:
      - file: Base64 encoded file data
      - inputFormat: Input format (pdf, docx, html, md, etc.)
      - outputFormat: Output format (pdf, docx, html, md, etc.)
    """
    file_data = params.get("file", "")
    input_format = params.get("inputFormat", "").lower()
    output_format = params.get("outputFormat", "").lower()
    
    if not file_data or not input_format or not output_format:
        await manager.send_error(message_id, 400, "File data, input format, and output format are required", websocket)
        return
    
    try:
        # Create temporary files for input and output
        with tempfile.NamedTemporaryFile(suffix=f".{input_format}", delete=False) as input_file:
            input_path = input_file.name
            input_file.write(base64.b64decode(file_data))
        
        output_file = tempfile.NamedTemporaryFile(suffix=f".{output_format}", delete=False)
        output_path = output_file.name
        output_file.close()
        
        # Use pandoc for conversion
        process = await asyncio.create_subprocess_shell(
            f"pandoc -f {input_format} -t {output_format} -o {output_path} {input_path}",
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE
        )
        stdout, stderr = await process.communicate()
        
        if process.returncode == 0:
            # Read the converted file
            with open(output_path, 'rb') as file:
                converted_data = base64.b64encode(file.read()).decode()
                result = {
                    "file": converted_data,
                    "format": output_format
                }
        else:
            # If pandoc failed, try to provide helpful error messages
            error_message = stderr.decode() if stderr else "Unknown error during conversion"
            result = {
                "error": error_message,
                "suggestion": "Make sure pandoc is installed and the formats are supported"
            }
        
        # Clean up temporary files
        try:
            os.unlink(input_path)
            os.unlink(output_path)
        except:
            pass
        
        await manager.send_response(message_id, result, websocket)
    
    except Exception as e:
        await manager.send_error(message_id, 500, str(e), websocket)

# Register document management methods
DOCUMENT_HANDLERS = {
    "parseDocument": handle_parse_document,
    "createDocument": handle_create_document,
    "convertDocument": handle_convert_document,
}

# Update the METHOD_HANDLERS dictionary
from src.main import METHOD_HANDLERS
METHOD_HANDLERS.update(DOCUMENT_HANDLERS)

# Make this runnable as a standalone server too
if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8002))
    host = os.environ.get("HOST", "0.0.0.0")
    print(f"Starting Document Management MCP Server on {host}:{port}")
    print(f"Available methods: {', '.join(DOCUMENT_HANDLERS.keys())}")
    uvicorn.run(app, host=host, port=port)